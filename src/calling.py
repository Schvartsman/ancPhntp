#!/usr/bin/env python3

import os
import argparse
import numpy as np
from docx import Document, document, table
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def argument():
    parser=argparse.ArgumentParser(description="Calling of ancient DNA",
                                   formatter_class=argparse.ArgumentDefaultsHelpFormatter)
    parser.add_argument("--sam", type=str,
                        help="Input sam-file for calling")
    parser.add_argument("--id", type=str,
                        help="ID of specimen")
    parser.add_argument("--ref", type=str,
                        help="Number of reference genome (19, 38c or 38n), which was applied in alignment")
    args=parser.parse_args()
    return args

def cigar(a):
    x=0
    y=0
    n=''
    l=''
    for i in range(len(a)):
        if a[i] != 'M' and a[i] != 'I' and a[i] != 'D' and a[i] != 'N' and a[i] != 'S' and a[i] != 'H' and a[i] != 'P' and a[i] != '=' and a[i] != 'X':
            n = n + a[i]
        else:
            l = a[i]
            if l == 'M':
                x = x + int(n)
                y = y + int(n)
            if l == 'I':
                x = x + int(n)
            if l == 'D':
                y = y + int(n)
            if l == 'N':
                y = y + int(n)
            if l == 'S':
                x = x + int(n)
            if l == '=':
                x = x + int(n)
                y = y + int(n)
            if l == 'X':
                x = x + int(n)
                y = y + int(n)
            l = ''
            n = ''
    return x,y

def regain(z,k,coor_o):
    if coor_o == 'A':
        z[k][0] = z[k][0] + 1
    if coor_o == 'C':
        z[k][1] = z[k][1] + 1
    if coor_o == 'G':
        z[k][2] = z[k][2] + 1
    if coor_o == 'T':
        z[k][3] = z[k][3] + 1
    return z

def coor(a,b,c,d):
    nuc=''
    x=-1
    y=-1
    n=''
    l=''
    for i in range(len(a)):
        if a[i] != 'M' and a[i] != 'I' and a[i] != 'D' and a[i] != 'N' and a[i] != 'S' and a[i] != 'H' and a[i] != 'P' and a[i] != '=' and a[i] != 'X':
            n = n + a[i]
        else:
            l = a[i]
            if l == 'M':
                if b+y+int(n) > c:
                    for j in range(int(n)):
                        y = y + 1
                        x = x + 1
                        if b+y == c:
                            nuc = d[x]
                    break
                if b+y+int(n) == c:
                    x = x + int(n)
                    y = y + int(n)
                    nuc = d[x]
                    break
                x = x + int(n)
                y = y + int(n)
            if l == 'I':
                x = x + int(n)
            if l == 'D':
                y = y + int(n)
            if l == 'N':
                y = y + int(n)
            if l == 'S':
                x = x + int(n)
            if l == '=':
                if b+y+int(n) > c:
                    for j in range(int(n)):
                        y = y + 1
                        x = x + 1
                        if b+y == c:
                            nuc = d[x]
                    break
                if b+y+int(n) == c:
                    x = x + int(n)
                    y = y + int(n)
                    nuc = d[x]
                    break
                x = x + int(n)
                y = y + int(n)
            if l == 'X':
                if b+y+int(n) > c:
                    for j in range(int(n)):
                        y = y + 1
                        x = x + 1
                        if b+y == c:
                            nuc = d[x]
                    break
                if b+y+int(n) == c:
                    x = x + int(n)
                    y = y + int(n)
                    nuc = d[x]
                    break
                x = x + int(n)
                y = y + int(n)
            l = ''
            n = ''
    return nuc

def code(a,b,c,d):
    res1 = ''
    res2 = ''
    res3 = ''
    for i in range(len(c)):
        y = c[i].split('\t')
        if y[0] == a:
            if b[i][0] != 0:
                res1 = res1 + 'A'
                res2 = res2 + str(b[i][0])[:-2]
            if b[i][1] != 0:
                res1, res2 = slash(res1,res2)
                res1 = res1 + 'C'
                res2 = res2 + str(b[i][1])[:-2]
            if b[i][2] != 0:
                res1, res2 = slash(res1,res2)
                res1 = res1 + 'G'
                res2 = res2 + str(b[i][2])[:-2]
            if b[i][3] != 0:
                res1, res2 = slash(res1,res2)
                res1 = res1 + 'T'
                res2 = res2 + str(b[i][3])[:-2]
            break
    if res1 == '':
        return 'NA', 'NA'
    else:
        if len(res1) == 1:
            if res1 == d:
                res3 = '2'
            else:
                res3 = '0'
        else:
            n = int((len(res1)+1)/2)
            check = False
            for i in range(n):
                if res1[i*2] == d:
                    check = True
                    k = i
            if check == False:
                res3 = 0
            else:
                array = np.zeros(n)
                num = ''
                for i in range(len(res2)):
                    if res2[i] == '|':
                        for j in range(n):
                            if array[j] == 0:
                                array[j] = int(num)
                                num = ''
                                break
                    else:
                        num = num + res2[i]
                gen = 0
                for i in range(n):
                    gen = gen + array[i]
                w = 0
                if n > 2:
                    for i in range(len(array)):
                        if array[i] == max(array):
                            w = w + 1
                    if w > 2:
                        res3 = 1
                    if array[k]/gen >= 0.65:
                        res3 = 2
                    if array[k]/gen < 0.65 and array[k]/gen >= 0.4:
                        res3 = 1
                    if array[k]/gen < 0.4:
                        res3 = 0
                if n == 2:
                    if array[k]/gen >= 0.65:
                        res3 = 2
                    if array[k]/gen < 0.65 and array[k]/gen >= 0.4:
                        res3 = 1
                    if array[k]/gen < 0.4:
                        res3 = 0
        return res1+':'+res2, res3

def slash(a,b):
    if len(a) != 0:
        a = a + '|'
        b = b + '|'
    return a,b

def ch(a):
    num = ''
    for i in range(len(a)-3):
        num = num + a[i+3]
        if a[i+3] != '0' and a[i+3] != '1' and a[i+3] != '2' and a[i+3] != '3' and a[i+3] != '4' and a[i+3] != '5' and a[i+3] != '6' and a[i+3] != '7' and a[i+3] != '8' and a[i+3] != '9':
            num = '0'
            break
    return int(num)

args = argument()
sam = args.sam
idd = args.id
ref = args.ref
tsv19 = '/home/daniils/tools/ancPhntp/data/19.tsv'
tsv38.c = '/home/daniils/tools/ancPhntp/data/38.chr.tsv'
tsv38.n = '/home/daniils/tools/ancPhntp/data/38.nc.tsv'
tsvw = '/home/daniils/tools/ancPhntp/data/specimen.tsv'
if ref == '19':
    t = open(tsv19, 'r')
if ref == '38c':
    t = open(tsv38.c, 'r')
if ref == '38n':
    t = open(tsv38.n, 'r')
times = t.readlines()
t.close()
f = open(sam, 'r')
k = 0
z = np.zeros([41,4])
y = times[k].split('\t')
kind = False
for s in f:
    if s[0] == '@':
        continue
    else:
        x = s.split('\t')
        if x[2] == '*':
            break
        while ch(y[1]) < ch(x[2]):
            k = k + 1
            y = times[k].split('\t')
        if x[2] == y[1]:
            p = 0
            if k < 40:
                while cigar(x[5])[1]+int(x[3]) >= int(y[2]) and int(x[3]) <= int(y[2]) and x[2] == y[1]:
                    z = regain(z,k+p,coor(x[5],int(x[3]),int(y[2]),x[9]))
                    p = p + 1
                    if p+k > 40:
                        break
                    y = times[k+p].split('\t')
            if k == 40:
                if cigar(x[5])[1]+int(x[3]) >= int(y[2]) and int(x[3]) <= int(y[2]):
                    z = regain(z,k+p,coor(x[5],int(x[3]),int(y[2]),x[9]))
            y = times[k].split('\t')
            while int(x[3]) > int(y[2]) and k < 40:
                if x[2] == y[1]:
                    if k < 40:
                        k = k + 1
                        y = times[k].split('\t')
                        p = 0
                        while cigar(x[5])[1]+int(x[3]) >= int(y[2]) and int(x[3]) <= int(y[2]) and x[2] == y[1]:
                            z = regain(z,k+p,coor(x[5],int(x[3]),int(y[2]),x[9]))
                            p = p + 1
                            if p+k > 40:
                                break
                            y = times[k+p].split('\t')
                        y = times[k].split('\t')
                    else:
                        y = times[k].split('\t')
                        if cigar(x[5])[1]+int(x[3]) >= int(y[2]) and int(x[3]) <= int(y[2]):
                            z = regain(z,k,coor(x[5],int(x[3]),int(y[2]),x[9]))
                else:
                    break
            if int(x[3]) > int(y[2]) and k == 40:
                break
f.close()
w = open(tsvw, 'r')
wines = w.readlines()
name = []
items = [[] for i in range(41)] 
for i in range(len(wines)):
    wew = wines[i].split('\t')
    if i == 0:
        for j in range(len(wew)):
            name.append(wew[j])
    if i != 0:
        if i != 36:
            for j in range(len(wew)+2):
                if j < 5:
                    items[i-1].append(wew[j])
                if j == 5:
                    items[i-1].append(wew[j][0])
                if j == 6:
                    items[i-1].append(code(wew[1],z,times,wew[5][0])[0])
                if j == 7:
                    items[i-1].append(code(wew[1],z,times,wew[5][0])[1])
        else:
            for j in range(len(wew)+2):
                if j < 5:
                    items[i-1].append(wew[j])
                if j == 5:
                    items[i-1].append(wew[j][0])
                if j == 6:
                    point36 = np.zeros(4)
                    letter = []
                    first_part = ''
                    second_part = ''
                    sign = True
                    word = code(wew[1],z,times,wew[5][0])[0]
                    for k in range(len(word)):
                        if sign == False:
                            second_part = second_part + word[k]
                        if word[k] == ':':
                            sign = False
                        if sign == True:
                            first_part = first_part + word[k]
                    for k in range(len(first_part)):
                        if first_part[k] != '|':
                            letter.append(first_part[k])
                    num = ''
                    num_po = 0
                    second_part = second_part+'|'
                    for k in range(len(second_part)):
                        if second_part[k] != '|':
                            num = num + second_part[k]
                        else:
                            if first_part[2*num_po] == 'A':
                                point36[0] = int(num)
                            if first_part[2*num_po] == 'C':
                                point36[1] = int(num)
                            if first_part[2*num_po] == 'G':
                                point36[2] = int(num)
                            if first_part[2*num_po] == 'T':
                                point36[3] = int(num)
                            num_po = num_po + 1
                            num = ''
                    gen = 0
                    for k in range(4):
                        gen = gen + point36[k]
                    items[i-1].append(code(wew[1],z,times,wew[5][0])[0])
                    print(point36, '  ', gen)
                if j == 7:
                    if gen != 0:
                        if point36[0]/gen > 0.4 and point36[2]/gen > 0.4:
                            items[i-1].append(code(wew[1],z,times,wew[5][0])[1])
                            continue
                        if point36[1]/gen > 0.4 and point36[3]/gen > 0.4:
                            items[i-1].append(code(wew[1],z,times,wew[5][0])[1])
                            continue
                        if point36[2] < 0.65 and point36[2] > 0.4:
                            items[i-1].append('1')
                        if point36[2] > 0.65:
                            items[i-1].append('2')
                        if point36[2] <= 0.4:
                            items[i-1].append('0')
                    if gen == 0:
                        items[i-1].append('NA')

doc = Document()
style = doc.styles['Normal']
style.font.name = 'TimesNewRoman'
style.font.size = Pt(8)
table = doc.add_table(1, len(items[0]))
table.autofit = False
for row in table.rows:
        row.cells[0].width = Inches(0.1)
table.style = 'Light Shading Accent 1'
head_cells = table.rows[0].cells
for i,item in enumerate(name):
    p = head_cells[i].paragraphs[0]
    p.add_run(item).bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for row in items:
    cells=table.add_row().cells
    for i, item in enumerate(row):
        cells[i].text=str(item)
doc.save('/home/daniils/'+idd+'.docx')
os.system("cp /home/daniils/tools/ancPhntp/data/phenotype.csv /home/daniils/"+idd+".csv")
t = open("/home/daniils/"+idd+".csv", 'a')
first = ''
for i in range(41):
    first = first + str(items[i][7])
    if i != 40:
        first = first + ','
t.write(first)
t.close()
